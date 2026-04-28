from __future__ import annotations

import csv
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parent
REPORT_TXT = ROOT / "report.txt"
DOCX_PATH = ROOT / "MID-COURSE PROJECT.docx"
SEQUENCE_METRICS_CSV = ROOT / "output" / "metrics" / "sequence_metrics.csv"
SUMMARY_TXT = ROOT / "output" / "metrics" / "summary.txt"


def is_heading(line: str) -> bool:
    stripped = line.strip()
    if not stripped:
        return False
    if stripped.startswith("- "):
        return False

    letters = [ch for ch in stripped if ch.isalpha()]
    if not letters:
        return False

    return all(ch.isupper() or not ch.isalpha() for ch in stripped)


def read_report_lines() -> list[str]:
    return REPORT_TXT.read_text(encoding="utf-8").splitlines()


def read_metrics() -> list[dict[str, str]]:
    if not SEQUENCE_METRICS_CSV.exists():
        return []

    with SEQUENCE_METRICS_CSV.open(newline="", encoding="utf-8") as handle:
        metrics = list(csv.DictReader(handle))

    order = {"bird": 0, "car": 1, "frog": 2, "sheep": 3, "squirrel": 4}
    metrics.sort(key=lambda item: order.get(item["category"], 999))
    return metrics


def read_summary() -> dict[str, str]:
    summary: dict[str, str] = {}
    if not SUMMARY_TXT.exists():
        return summary

    for line in SUMMARY_TXT.read_text(encoding="utf-8").splitlines():
        if ":" not in line:
            continue
        key, value = line.split(":", 1)
        summary[key.strip()] = value.strip()
    return summary


def set_page_layout(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)

    for style_name, size in [("Heading 1", 14), ("Heading 2", 12)]:
        style = document.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)


def add_text_content(document: Document, lines: list[str]) -> None:
    non_empty = [line for line in lines if line.strip()]
    if len(non_empty) < 2:
        return

    title = non_empty[0]
    subtitle = non_empty[1]

    title_p = document.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run(title)
    run.bold = True
    run.font.size = Pt(16)

    subtitle_p = document.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_p.add_run(subtitle).bold = True

    document.add_paragraph()

    content_lines = lines[lines.index(subtitle) + 1 :]
    paragraph_buffer: list[str] = []

    def flush_paragraph() -> None:
        if not paragraph_buffer:
            return
        text = " ".join(part.strip() for part in paragraph_buffer)
        document.add_paragraph(text)
        paragraph_buffer.clear()

    for raw_line in content_lines:
        line = raw_line.rstrip()

        if not line.strip():
            flush_paragraph()
            continue

        if is_heading(line):
            flush_paragraph()
            document.add_paragraph(line.strip(), style="Heading 2")
            continue

        if line.startswith("- "):
            flush_paragraph()
            document.add_paragraph(line[2:].strip(), style="List Bullet")
            continue

        paragraph_buffer.append(line)

    flush_paragraph()


def add_metrics_table(document: Document, metrics: list[dict[str, str]], summary: dict[str, str]) -> None:
    if not metrics:
        return

    document.add_paragraph("Performance Table", style="Heading 2")
    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"

    hdr = table.rows[0].cells
    hdr[0].text = "Category"
    hdr[1].text = "Predicted Box"
    hdr[2].text = "Ground Truth Box"
    hdr[3].text = "IoU"

    for metric in metrics:
        row = table.add_row().cells
        row[0].text = metric["category"]
        row[1].text = (
            f'{metric["pred_xmin"]} {metric["pred_ymin"]} '
            f'{metric["pred_xmax"]} {metric["pred_ymax"]}'
        )
        row[2].text = (
            f'{metric["label_xmin"]} {metric["label_ymin"]} '
            f'{metric["label_xmax"]} {metric["label_ymax"]}'
        )
        row[3].text = metric["iou"]

    if summary:
        row = table.add_row().cells
        row[0].text = "Mean IoU"
        row[1].text = ""
        row[2].text = ""
        row[3].text = summary.get("mean_iou", "")

        row = table.add_row().cells
        row[0].text = "Detection Accuracy"
        row[1].text = ""
        row[2].text = ""
        row[3].text = summary.get("detection_accuracy", "")


def add_output_images(document: Document, metrics: list[dict[str, str]]) -> None:
    if not metrics:
        return

    document.add_page_break()
    document.add_paragraph("Output Images", style="Heading 2")

    for index, metric in enumerate(metrics, start=1):
        category = metric["category"]
        image_path = ROOT / "output" / "results" / category / "0000.png"
        if not image_path.exists():
            continue

        image_paragraph = document.add_paragraph()
        image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        image_paragraph.add_run().add_picture(str(image_path), width=Inches(5.8))

        caption = document.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.add_run(
            f"Figure {index}. {category.capitalize()} sequence output. "
            f'Predicted IoU: {metric["iou"]}.'
        ).italic = True


def main() -> None:
    lines = read_report_lines()
    metrics = read_metrics()
    summary = read_summary()

    document = Document()
    set_page_layout(document)
    add_text_content(document, lines)
    add_metrics_table(document, metrics, summary)
    add_output_images(document, metrics)
    document.save(DOCX_PATH)


if __name__ == "__main__":
    main()
